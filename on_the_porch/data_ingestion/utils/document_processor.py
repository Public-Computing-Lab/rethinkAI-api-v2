"""
Document processing utilities for extracting text from various file formats.
"""
from pathlib import Path
from typing import List, Dict, Any
from datetime import datetime

from pypdf import PdfReader
import docx
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document


def extract_text_from_pdf(file_path: Path) -> str:
    """Extract text from a PDF file (all pages combined)."""
    try:
        reader = PdfReader(file_path)
        text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
        return text.strip()
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {e}")


def extract_pages_from_pdf(file_path: Path) -> List[Dict[str, Any]]:
    """
    Extract text from a PDF file page by page.
    
    Returns:
        List of dicts with 'page_num' and 'text' for each page
    """
    try:
        reader = PdfReader(file_path)
        pages = []
        for i, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                pages.append({
                    'page_num': i,
                    'text': text,
                    'total_pages': len(reader.pages)
                })
        return pages
    except Exception as e:
        raise ValueError(f"Failed to extract pages from PDF: {e}")


def extract_text_from_docx(file_path: Path) -> str:
    """Extract text from a DOCX file."""
    try:
        doc = docx.Document(file_path)
        paragraphs = [paragraph.text for paragraph in doc.paragraphs]
        return "\n\n".join(paragraphs).strip()
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {e}")


def extract_text_from_txt(file_path: Path) -> str:
    """Extract text from a plain text file."""
    try:
        return file_path.read_text(encoding='utf-8', errors='ignore').strip()
    except Exception as e:
        raise ValueError(f"Failed to read text file: {e}")


def extract_text_from_file(file_path: Path) -> str:
    """
    Extract text from various file formats.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If file format is unsupported or extraction fails
    """
    ext = file_path.suffix.lower()
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext == '.doc':
        # For old .doc files, try DOCX library (may not work for all)
        return extract_text_from_docx(file_path)
    elif ext in {'.txt', '.md'}:
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# Chunk size settings per folder category
CHUNK_SETTINGS = {
    'newsletters': {'chunk_size': 2500, 'chunk_overlap': 300, 'page_wise': True},
    'policy': {'chunk_size': 1500, 'chunk_overlap': 200, 'page_wise': False},
    'transcripts': {'chunk_size': 1500, 'chunk_overlap': 200, 'page_wise': False},
    'default': {'chunk_size': 1000, 'chunk_overlap': 200, 'page_wise': False},
}


def process_file_to_documents(
    file_path: Path,
    file_metadata: Dict[str, Any],
    chunk_size: int = None,
    chunk_overlap: int = None
) -> List[Document]:
    """
    Process a file into LangChain documents for the vector DB.
    
    Args:
        file_path: Path to the file
        file_metadata: Metadata about the file (id, name, modifiedTime, etc.)
        chunk_size: Size of text chunks (auto-determined by folder_category if None)
        chunk_overlap: Overlap between chunks (auto-determined by folder_category if None)
        
    Returns:
        List of Document objects ready for vector DB ingestion
    """
    folder_category = file_metadata.get('folder_category', 'root')
    settings = CHUNK_SETTINGS.get(folder_category, CHUNK_SETTINGS['default'])
    
    # Use provided values or fall back to settings
    chunk_size = chunk_size or settings['chunk_size']
    chunk_overlap = chunk_overlap or settings['chunk_overlap']
    use_page_wise = settings.get('page_wise', False)
    
    ext = file_path.suffix.lower()
    
    # For newsletters PDFs, process page by page
    if use_page_wise and ext == '.pdf':
        return _process_pdf_page_wise(file_path, file_metadata, chunk_size, chunk_overlap)
    
    # Standard processing for other files
    text = extract_text_from_file(file_path)
    
    if not text.strip():
        return []
    
    # Split into chunks
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    chunks = text_splitter.split_text(text)
    
    # Decide doc_type based on folder category (for compatibility with older RAG)
    fc_lower = folder_category.lower()
    if fc_lower == 'policy':
        doc_type = 'policy'
    elif fc_lower in ('transcript', 'transcripts'):
        doc_type = 'transcript'
    else:
        doc_type = 'client_upload'

    # Create documents with metadata
    documents = []
    for i, chunk in enumerate(chunks):
        doc = Document(
            page_content=chunk,
            metadata={
                'source': file_metadata.get('name', 'unknown'),
                'doc_type': doc_type,
                'folder_category': folder_category,
                'chunk_id': i,
                'drive_file_id': file_metadata.get('id', ''),
                'modified_time': file_metadata.get('modifiedTime', ''),
                'ingestion_date': datetime.now().isoformat(),
                'file_extension': ext
            }
        )
        documents.append(doc)
    
    return documents


def _process_pdf_page_wise(
    file_path: Path,
    file_metadata: Dict[str, Any],
    chunk_size: int,
    chunk_overlap: int
) -> List[Document]:
    """
    Process a PDF file page by page, creating chunks for each page separately.
    This preserves page context and is better for newsletters.
    """
    folder_category = file_metadata.get('folder_category', 'root')
    pages = extract_pages_from_pdf(file_path)
    
    if not pages:
        return []
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    
    # Decide doc_type based on folder category (for compatibility with older RAG)
    fc_lower = folder_category.lower()
    if fc_lower == 'policy':
        doc_type = 'policy'
    elif fc_lower in ('transcript', 'transcripts'):
        doc_type = 'transcript'
    else:
        doc_type = 'client_upload'

    documents = []
    chunk_id = 0
    
    for page_info in pages:
        page_num = page_info['page_num']
        total_pages = page_info['total_pages']
        page_text = page_info['text']
        
        if not page_text.strip():
            continue
        
        # Split this page's text into chunks
        page_chunks = text_splitter.split_text(page_text)
        
        for chunk in page_chunks:
            doc = Document(
                page_content=chunk,
                metadata={
                    'source': file_metadata.get('name', 'unknown'),
                    'doc_type': doc_type,
                    'folder_category': folder_category,
                    'chunk_id': chunk_id,
                    'page_num': page_num,
                    'total_pages': total_pages,
                    'drive_file_id': file_metadata.get('id', ''),
                    'modified_time': file_metadata.get('modifiedTime', ''),
                    'ingestion_date': datetime.now().isoformat(),
                    'file_extension': '.pdf'
                }
            )
            documents.append(doc)
            chunk_id += 1
    
    return documents


def validate_file(file_path: Path, supported_extensions: set) -> bool:
    """
    Check if a file is valid and supported.
    
    Args:
        file_path: Path to the file
        supported_extensions: Set of supported file extensions (e.g., {'.pdf', '.docx'})
        
    Returns:
        True if file is valid, False otherwise
    """
    if not file_path.exists():
        return False
    
    if not file_path.is_file():
        return False
    
    if file_path.suffix.lower() not in supported_extensions:
        return False
    
    # Check file size (skip very large files > 50MB)
    if file_path.stat().st_size > 50 * 1024 * 1024:
        return False
    
    return True


def get_file_info(file_path: Path) -> Dict[str, Any]:
    """
    Get basic information about a file.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Dictionary with file information
    """
    stat = file_path.stat()
    return {
        'name': file_path.name,
        'size': stat.st_size,
        'modified': datetime.fromtimestamp(stat.st_mtime).isoformat(),
        'extension': file_path.suffix.lower()
    }

